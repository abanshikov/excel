import pandas as pd
import docx
from modules.json_files import JSON


EMPTY_CELLS = [0, '???', '?', '-', ' -', ' Нет', 'Нет', 'Не знаю', 'нет',  'Никогда', 'Неразборчиво', '', ]
MILITARY_SERVICE = ['Не служил', 'Не служил, Турецкий подданый', 'Не служил, освобожден по льготе']
EDUCATION = ['Не обучался', 'Нет', 'Неразборчиво', 'Не помнит', '',]

df = pd.read_excel('data/data.xls', header=None, skiprows=[0, 1, 2, 3])


def get_declination(age: int) -> str:
    if (age % 10 == 1) and (age != 11) and (age != 111):
        return "год"
    elif (age % 10 > 1) and (age % 10 < 5) and (age != 12) and (age != 13) and (age != 14):
        return "года"
    else:
        return "лет"


def lower_register(my_str) -> str:
    """
    Уменьшение первой буквы строки
    :param my_str:
    :return:
    """
    return " ".join([word[0].lower() + word[1:] for word in str(my_str).strip().split()])


def get_info_block(rows: pd.Series) -> str:
    """
    Получение данных для информационного блока
    :param rows:
    :return:
    """
    # ФИО
    try:
        text = f'{rows[4].strip()} '
    except:
        text = ''

    # Дата рождения
    if rows[5] not in EMPTY_CELLS:
        text += f'{str(int(rows[5]))} года рождения, '
    text += 'крымчак, '

    # Профессия
    if rows[14] not in EMPTY_CELLS:
        text += f'{lower_register(rows[14])}, '

    # Сословие
    if rows[7] not in EMPTY_CELLS:
        text += f'{lower_register(rows[7])}, '

    # Сообщество
    if rows[15] not in EMPTY_CELLS:
        if rows[15] == 'Да':
            text += 'состоял членом обществ, '
        elif rows[15] == 'Сведения не пожелал дать':
            text += 'сведения не пожелал дать, '
        else:
            text += f'членство в обществах: {str(rows[15]).strip().replace("Член", "член")}, '

    # Военная служба
    if str(rows[16]).strip() not in EMPTY_CELLS+MILITARY_SERVICE and 'Не служил' not in str(rows[16]).strip():
        text += f'на военной службе состоял, '
    else:
        text += f'на военной службе не состоял, '

    # Владение грамотой по русски
    if str(rows[9]).strip() == 'Грамотный':
        text += f'по-русски грамотен, '
    elif str(rows[9]).strip() == 'Малограмотный':
        text += f'по-русски малограмотен, '
    else:
        text += f'по-русски не грамотен, '

    # Владение грамотой по др.-евр.
    if str(rows[8]).strip() == 'Грамотный':
        text += f'по др.-евр. образован, '
    elif str(rows[8]).strip() == 'Малограмотный':
        text += f'по др.-евр. малообразован, '
    else:
        text += f'по др.-евр. не образован, '

    # Семейное положение
    if rows[17] not in EMPTY_CELLS:
        text += f'{str(rows[17]).strip().lower()}.'

    return text


def get_info_wife_block(rows) -> str:
    """
    Получение данных для информационного блока
    :param rows:
    :return:
    """
    # ФИО
    try:
        text = f'{rows[21].strip()} '
    except:
        text = ''

    # Дата рождения
    if rows[22] not in EMPTY_CELLS:
        text += f'{str(int(rows[22]))} года рождения, '
    text += 'крымчачка, '

    # Профессия
    if rows[31] not in EMPTY_CELLS:
        text += f'{lower_register(rows[31])}, '

    # Сословие
    if rows[24] not in EMPTY_CELLS:
        text += f'{lower_register(rows[24])}, '

    # Сообщество
    if rows[15] not in EMPTY_CELLS:
        if rows[15] == 'Да':
            text += 'состояла членом обществ, '
        elif rows[15] == 'Сведения не пожелал дать':
            text += 'сведения не пожелал дать, '
        else:
            text += f'членство в обществах: {str(rows[15]).strip().replace("Член", "член")}, '

    # Грамотность жены по-русски
    if str(rows[26]).strip() == 'Грамотная':
        text += f' грамотная по-русски,'
    elif str(rows[26]).strip() == 'Малограмотная':
        text += f' малограмотная по-русски,'
    else:
        text += f' не грамотная по-русски,'

    # Грамотность жены по д-ев
    if str(rows[25]).strip() == 'Грамотная':
        text += f' грамотная по др.-евр.'
    elif str(rows[25]).strip() == 'Малограмотная':
        text += f' малограмотная по др.-евр.'
    else:
        text += f' не грамотная по др.-евр.'

    # Образование жены по русски
    if rows[28] not in EMPTY_CELLS + EDUCATION:
        text += f', обучалась по русски: {str(rows[28]).strip()}'

    if rows[30] not in EMPTY_CELLS:
        if str(rows[30]).strip().isdigit():
            text += f' {str(rows[30]).strip()} {get_declination(int(rows[30]))}'
        else:
            text += f' {str(rows[30]).strip()}'

    # Образование по др.-евр. языкам
    if rows[27] not in EMPTY_CELLS + EDUCATION:
        text += f'; обучалась по др.-евр.: {str(rows[27]).strip()}'

    if rows[29] not in EMPTY_CELLS:
        if str(rows[29]).strip().isdigit():
            text += f' {str(rows[29]).strip()} {get_declination(int(rows[29]))}.'
        else:
            text += f' {str(rows[29]).strip()}.'
    else:
        text += '.'

    # Семейное положение
    if rows[32] not in EMPTY_CELLS:
        text += f' {str(rows[32]).strip().capitalize()}.'

    return text


def get_press(rows) -> str:
    text = ''
    # Чтение газет
    if rows[45] not in EMPTY_CELLS:
        text += f' {str(rows[45]).strip()},'
    else:
        text += f' Ни газет, ни журналов не читает,'

    # Имеет ли абонементы в  библиотеке
    if rows[46] not in EMPTY_CELLS:
        text += f' в библиотеке абонировался,'
    else:
        text += f' в библиотеке не абонировался,'

    # Выписывает  ли   журналы, газеты и книги
    if rows[47] not in EMPTY_CELLS and str(rows[47]).strip() != 'Да':
        text += f' {lower_register(rows[47])},'
    elif str(rows[47]).strip() != 'Да':
        text += ' периодические издания (СМИ: газеты, журналы и пр.) и книги выписывал,'
    else:
        text += f' периодические издания (СМИ: газеты, журналы и пр.) и книги не выписывал,'

    # Посещает ли клуб или какие-либо общественные собрания
    if rows[48] not in EMPTY_CELLS:
        text += f' {lower_register(rows[48])},'
    else:
        text += f' клуб и общественные собрания не посещал,'

    # Слушает ли лекции
    if rows[49] not in EMPTY_CELLS:
        text += f' слушал лекции по распространению политических и научных знаний,'
    else:
        text += f' лекций по распространению политических и научных знаний не слушал,'

    # Часто ли бывает в театре
    if rows[50] not in EMPTY_CELLS and str(rows[50]).strip() != 'Никогда не ходил':
        text += f' в театр ходил {lower_register(rows[50])},'
    else:
        text += f' в театр никогда не ходил,'

    # Посещает ли иллюзион
    if rows[51] not in EMPTY_CELLS:
        text += f' посещал иллюзион.'
    else:
        text += f' иллюзион не посещал.'

    # Знаком ли с преданиями о происхождении крымчаков
    if rows[53] == 'Да':
        text += f' С преданиями о происхождении крымчаков знаком.'
    else:
        text += f' С преданиями о происхождении крымчаков не знаком.'

    # Посещает  ли  синагогу   в будни
    if rows[52] not in EMPTY_CELLS and str(rows[52]).strip() != 'Нет, по субботам':
        text += f' Синагогу в будни посещал. '
    else:
        text += f' Синагогу в будни не посещал. '

    return text


def get_education_block(rows):
    """
    Получение данных для блока образования
    :param rows:
    :return:
    """
    if rows[10] not in EMPTY_CELLS + EDUCATION or rows[11] not in EMPTY_CELLS+EDUCATION:
        text = 'В детстве обучался'
    else:
        text = 'В детстве не обучался'

    # Образование по русски
    if rows[11] not in EMPTY_CELLS+EDUCATION:
        text += f' по русски: {str(rows[11]).strip()},'
    else:
        text += ','

    if rows[13] not in EMPTY_CELLS:
        if str(rows[13]).strip().isdigit():
            text += f' {str(rows[13]).strip()} {get_declination(int(rows[13]))} изучал русский и'
        else:
            text += f' {str(rows[13]).strip()} лет изучал русский и'

    # Образование по др.-евр. языкам
    if rows[12] not in EMPTY_CELLS or rows[10] not in EMPTY_CELLS+EDUCATION:
        if rows[12] not in EMPTY_CELLS:
            if str(rows[12]).strip().isdigit():
                text += f' {str(rows[12]).strip()} {get_declination(int(rows[12]))}'
            else:
                text += f' {str(rows[12]).strip()} лет'

        if rows[10] not in EMPTY_CELLS+EDUCATION:
            text += f' у {str(rows[10]).strip()}'
        text += ' др.-евр. языкам;'

    # Владение грамотой по русски
    if str(rows[9]).strip() == 'Грамотный':
        text += f'по-русски грамотен, '
    elif str(rows[9]).strip() == 'Малограмотный':
        text += f'по-русски малограмотен, '
    else:
        text += f'по-русски не грамотен, '

    # Владение грамотой по др.-евр.
    if str(rows[8]).strip() == 'Грамотный':
        text += f'по др.-евр. образован, '
    elif str(rows[8]).strip() == 'Малограмотный':
        text += f'по др.-евр. малообразован, '
    else:
        text += f'по др.-евр. не образован, '

    # Владение языками
    if rows[64] not in EMPTY_CELLS:
        text += f'разговаривал на языке(ах): {str(rows[64]).strip()}.'

    text += get_press(rows)

    return text


def get_socio_economic_block(rows: pd.Series) -> str:
    """
    Получение данных для социально-экономического блока
    :param rows:
    :return:
    """
    # В своем ли помещении работает, торгует
    other = ['Служит', 'Служит в чужом магазине', 'Служит в Синагоге', 'Торгует на базаре', 'Гастролирует',
             'Служащий',]
    premises = ['Не в своем помещении у отца', 'Не в своем помещении', 'Да, в своем помещении',
                'Не в своем помещении у хозяина', 'Не в своем помещении, мастерская отца',
                'Не в своем помещении у хозяина приказчика', 'Не в своем помещении с отцом',
                'Да в своем помещении, на дому', 'Да в своем помещении с отцом', 'Старый Крым',
                'Не в своем помещении у брата', 'Иногда в наемном', 'При отце А.Я. Хондо',
                'При отце А.Я.Хондо', 'Не в своем помещении, фабрика халвы', 'Не в своем помещении, работает у отца',
                 'Не в своем, служит в Крымчакской Синагоге', 'В балагане', 'Да, поземельный налог', 'В лавке', ]

    if rows[35] not in EMPTY_CELLS:
        if rows[35] in other:
            text = f'{str(rows[35]).strip()},'
        elif rows[35] in premises:
            text = f'Работал/торговал {lower_register(rows[35]).replace("да, ", "").replace("да ", "")},'
        else:
            text = f'Данных о рабочем помещении нет,'
    else:
        text = f'Данных о рабочем помещении нет,'

    # Если помещение съемное, какова арендная плата, руб.
    if rows[36] not in EMPTY_CELLS+['Ничего', 'Неизвестно',]:
        text += f' наёмное помещение обходилось {str(rows[36]).strip()} рублей в год,'

    # Имеет ли недвижимость в собственности
    if rows[37] not in EMPTY_CELLS:
        text += f' имел собственность ({lower_register(rows[37]).replace("да, ", "").replace("да ", "")}),'
        # Доход от собственности
        if rows[38] not in EMPTY_CELLS+['Ничего', 'Неразборчиво']:
            if str(rows[38]).isdigit():
                text += f' которая приносила доход {lower_register(rows[37])} руб.'
            else:
                text += f' которая приносила {lower_register(rows[37]).replace("Приносит ", "")}.'
        else:
            text += ' которая ничего не приносила.'
    else:
        text += f' не имел собственности.'

    # Во  сколько  обходится обучение детей, руб.
    if rows[42] in ['Дети взрослые', 'Нет детей']:
        text += f' {rows[42]},'
    elif str(rows[42]).strip() == 'Бесплатно':
        text += f' Дети учились бесплатно,'
    elif str(rows[42]).strip() == 'На благотворительственном счете (бесплатно)':
        text += f' Дети учились бесплатно,'
    elif str(rows[42]).strip() == 'Не учатся':
        text += f' Дети не учатся,'
    elif str(rows[42]).strip().isdigit():
        text += f' Обучение детей обходилось в {str(rows[42]).strip()} руб.,'
    else:
        text += f' Информации о расходах на обучение детей нет,'

    # Сколько тратит на синагогальные нужды, руб.
    if str(rows[43]).strip().isdigit():
        text += f' на синагогальные нужды тратил {str(rows[43]).strip()} руб. в год.'
    else:
        text += f' данные о тратах на синагогальные нужды отсутствуют.'

    # В своем ли доме живет
    if rows[39] == 'В своем доме':
        text += f' Проживал в своем доме.'
    elif rows[39] not in EMPTY_CELLS:
        text += f' Проживал не в своем доме'
        # Если нет, то какова арендная плата, руб.
        if str(rows[40]).strip().isdigit():
            text += f', оплачивая арендную плату в размере {str(rows[40]).strip()} руб. в год.'
        else:
            text += f' данные о тратах на синагогальные нужды отсутствуют.'
    else:
        text += f' Данных о доме нет.'

    return text


def get_family_block(rows: pd.Series) -> str:
    """
    Получение данных для брачно-семейного блока
    :param rows:
    :return:
    """
    # Когда женился
    try:
        text = f'В {str(int(rows[18]))} году '
    except:
        return ''

    # Во сколько лет
    try:
        text += f'в возрасте {str(int(rows[19]))} {get_declination(int(rows[19]))} создал семью с'
    except:
        text += 'создал семью с '

    # Возраст и ФИО жены
    try:
        text += f' {str(int(rows[33]))} летней {str(rows[21]).strip()},'
    except:
        text += f' {str(rows[21]).strip()},'

    # Год рождения жены
    if rows[22] not in EMPTY_CELLS:
        text += f' {str(int(rows[22])).strip()} года рождения,'

    # Сословие жены
    if rows[24] not in EMPTY_CELLS:
        text += f' {lower_register(rows[24])},'

    # Грамотность жены по-русски
    if str(rows[26]).strip() == 'Грамотная':
        text += f' грамотная по-русски,'
    elif str(rows[26]).strip() == 'Малограмотная':
        text += f' малограмотная по-русски,'
    else:
        text += f' не грамотная по-русски,'

    # Грамотность жены по д-ев
    if str(rows[25]).strip() == 'Грамотная':
        text += f' грамотная по др.-евр.'
    elif str(rows[25]).strip() == 'Малограмотная':
        text += f' малограмотная по др.-евр.'
    else:
        text += f' не грамотная по др.-евр.'

    # Образование жены по русски
    if rows[28] not in EMPTY_CELLS + EDUCATION:
        text += f', обучалась по русски: {str(rows[28]).strip()}'

    if rows[30] not in EMPTY_CELLS:
        if str(rows[30]).strip().isdigit():
            text += f' {str(rows[30]).strip()} {get_declination(int(rows[30]))}'
        else:
            text += f' {str(rows[30]).strip()}'

    # Образование по др.-евр. языкам
    if rows[27] not in EMPTY_CELLS + EDUCATION:
        text += f'; обучалась по др.-евр.: {str(rows[27]).strip()}'

    if rows[29] not in EMPTY_CELLS:
        if str(rows[29]).strip().isdigit():
            text += f' {str(rows[29]).strip()} {get_declination(int(rows[29]))}.'
        else:
            text += f' {str(rows[29]).strip()}.'
    else:
        text += '.'

    # Как женился
    if str(rows[60]).strip() == 'По любви':
        text += f' Женился по любви,'
    elif str(rows[60]).strip() == 'По сватовству':
        text += f' Женился по сватовству,'
    else:
        text += f' Женился,'

    # Сколько приданого получил
    if rows[61] not in EMPTY_CELLS:
        if str(rows[61]).strip().isdigit():
            text += f' приданого за невесту получил {str(rows[61]).strip()} руб.,'
        else:
            text += f' приданого за невесту получил,'
    else:
        text += f' приданого за невесту не получил,'

    # Есть ли при семье старики, старухи, как зовут, занятие их и
    if rows[57] not in EMPTY_CELLS+['Не имеются', 'Есть', 'сам']:
        text += f' в семье проживали старики ({str(rows[57]).strip()}),'

    # Имеются ли в семье нетрудоспособные и с какими физическими недостатками
    if rows[56] not in EMPTY_CELLS+['Никто', 'Не имеем', 'Не имеются', 'Здоровые']:
        text += f' имелись нетрудоспособные (с физическими недостатками),'
    else:
        text += f' не имелись нетрудоспособные (с физическими недостатками),'

    # В каком возрасте умерли родители
    if str(rows[58]).lower() in ['жив', 'живы', 'жива', '']:
        text += f' отец жив,'
    elif rows[58] not in EMPTY_CELLS and str(rows[58]).strip().isdigit():
        text += f' отец умер в возрасте {str(rows[58]).strip()} {get_declination(int(rows[58]))},'
    else:
        text += f' отец умер,'

    if str(rows[59]).lower() in ['жив', 'живы', 'жива', '']:
        text += f' мать жива,'
    elif rows[59] not in EMPTY_CELLS and str(rows[59]).strip().isdigit():
        text += f' мать умерла в возрасте {str(rows[59]).strip()} {get_declination(int(rows[59]))},'
    else:
        text += f' мать умерла,'

    # Часто ли болеют члены семьи, и если определенной болезнью, то
    if rows[55] in ['Никто', 'Никогда']:
        text += f' члены семьи не болели.'
    elif rows[55] in ['Редко', 'очень редко ', '', '', '',]:
        text += f' члены семьи болели редко.'
    elif rows[55] not in EMPTY_CELLS+['Да']:
        text += f' члены семьи болели ({str(rows[55]).strip()}).'
    else:
        text += f' члены семьи болели.'

    # Сколько детей
    if rows[62] not in EMPTY_CELLS:
        text += f', В живых осталось {str(rows[62]).strip()} детей'
    if rows[63] not in EMPTY_CELLS:
        text += f', {str(rows[62]).strip()} умерли в детстве'
    text += '.'

    return text


def get_children_block(rows) -> str:
    """
    Получение данных для детского блока
    :param rows:
    :return:
    """
    sons = [
        {
            'name': 68,
            'age': 69,
            'last_study': 70,
            'current_study': 71,
            'business': 72,
            'jew': 73,
            'rus': 74,
            'free_time': 75,
        },
        {
            'name': 76,
            'age': 77,
            'last_study': 78,
            'current_study': 79,
            'business': 80,
            'jew': 81,
            'rus': 82,
            'free_time': 83,
        },
        {
            'name': 84,
            'age': 85,
            'last_study': 86,
            'current_study': 87,
            'business': 88,
            'jew': 89,
            'rus': 90,
            'free_time': 91,
        },
        {
            'name': 92,
            'age': 93,
            'last_study': 94,
            'current_study': 95,
            'business': 96,
            'jew': 97,
            'rus': 98,
            'free_time': 99,
        },
        {
            'name': 100,
            'age': 101,
            'last_study': 102,
            'current_study': 103,
            'business': 104,
            'jew': 105,
            'rus': 106,
            'free_time': 107,
        },
        {
            'name': 108,
            'age': 109,
            'last_study': 110,
            'current_study': 111,
            'business': 112,
            'jew': 113,
            'rus': 114,
            'free_time': 115,
        },
    ]
    daughters = [
        {
            'name': 117,
            'age': 118,
            'last_study': 119,
            'current_study': 120,
            'business': 121,
            'jew': 122,
            'rus': 123,
            'free_time': 124,
        },
        {
            'name': 125,
            'age': 126,
            'last_study': 127,
            'current_study': 128,
            'business': 129,
            'jew': 130,
            'rus': 131,
            'free_time': 132,
        },
        {
            'name': 133,
            'age': 134,
            'last_study': 135,
            'current_study': 136,
            'business': 137,
            'jew': 138,
            'rus': 139,
            'free_time': 140,
        },
        {
            'name': 141,
            'age': 142,
            'last_study': 143,
            'current_study': 144,
            'business': 145,
            'jew': 146,
            'rus': 147,
            'free_time': 148,
        },
        {
            'name': 149,
            'age': 150,
            'last_study': 151,
            'current_study': 152,
            'business': 153,
            'jew': 154,
            'rus': 155,
            'free_time': 156,
        },
        {
            'name': 157,
            'age': 158,
            'last_study': 159,
            'current_study': 160,
            'business': 161,
            'jew': 162,
            'rus': 163,
            'free_time': 164,
        },
        {
            'name': 165,
            'age': 166,
            'last_study': 167,
            'current_study': 168,
            'business': 169,
            'jew': 170,
            'rus': 171,
            'free_time': 172,
        },
    ]
    text_sons = ''
    text_daughters = ''

    # Сыновья
    for i, son in enumerate(sons):
        if rows[son["name"]] in EMPTY_CELLS:
            continue

        if i == 0:
            text_sons = 'Сын '

        # Имя
        if rows[son["name"]] not in EMPTY_CELLS:
            text_sons += rows[son["name"]]
        else:
            continue

        # Возраст
        if rows[son["age"]] not in EMPTY_CELLS:
            if str(rows[son["age"]]).strip().isdigit():
                text_sons += f' {str(rows[son["age"]]).strip()} лет от роду'
            else:
                text_sons += f' {str(rows[son["age"]]).strip()}'

        # Учился
        if rows[son["last_study"]] not in EMPTY_CELLS:
            text_sons += f', обучался в {str(rows[son["last_study"]]).strip()}'

        # Учится
        if rows[son["current_study"]] not in EMPTY_CELLS:
            text_sons += f', учится в {str(rows[son["current_study"]]).strip()}'

        # Сколько лет учился
        if rows[son["rus"]] not in EMPTY_CELLS:
            text_sons += f', изучал {str(rows[son["rus"]]).strip()} лет русский язык'
        if rows[son["jew"]] not in EMPTY_CELLS:
            text_sons += f', {str(rows[son["jew"]]).strip()} лет др.-евр. грамоту'

        # Занятие
        if rows[son["business"]] not in EMPTY_CELLS:
            text_sons += f', осваивал профессию {lower_register(rows[son["business"]])}'

        # Где проводит свободное время
        if rows[son["free_time"]] not in EMPTY_CELLS:
            text_sons += f', свободное время проводил {lower_register(rows[son["free_time"]])}'
        text_sons += '; '
    text_sons = text_sons[:-2] + '.'

    # Дочери
    for i, daughter in enumerate(daughters):
        if rows[daughter["name"]] in EMPTY_CELLS:
            continue

        if i == 0:
            text_daughters = 'Дочь '

        # Имя
        if rows[daughter["name"]] not in EMPTY_CELLS:
            text_daughters += rows[daughter["name"]]
        else:
            continue

        # Возраст
        if rows[daughter["age"]] not in EMPTY_CELLS:
            if str(rows[daughter["age"]]).strip().isdigit():
                text_daughters += f' {str(rows[daughter["age"]]).strip()} лет от роду'
            else:
                text_daughters += f' {str(rows[daughter["age"]]).strip()}'

        # Учился
        if rows[daughter["last_study"]] not in EMPTY_CELLS:
            text_daughters += f', училась в {str(rows[daughter["last_study"]]).strip()}'

        # Учится
        if rows[daughter["current_study"]] not in EMPTY_CELLS:
            text_daughters += f', учится в {str(rows[daughter["current_study"]]).strip()}'

        # Сколько лет учился
        if rows[daughter["rus"]] not in EMPTY_CELLS:
            text_daughters += f', {str(rows[daughter["rus"]]).strip()} лет изучала русский язык'
        if rows[daughter["jew"]] not in EMPTY_CELLS:
            text_daughters += f', {str(rows[daughter["jew"]]).strip()} лет обучалась др.-евр. грамоте'

        # Занятие
        if rows[daughter["business"]] not in EMPTY_CELLS:
            text_daughters += f', осваивала профессию {lower_register(rows[daughter["business"]])}'

        # Где проводит свободное время
        if rows[daughter["free_time"]] not in EMPTY_CELLS:
            text_daughters += f', свободное время проводила {lower_register(rows[daughter["free_time"]])}'
        text_daughters += '; '
    text_daughters = text_daughters[:-2] + '.'

    if text_sons == '.': text_sons = ''
    if text_daughters == '.': text_daughters = ''

    return text_sons, text_daughters


def create_text(rows: pd.Series) -> dict:
    """
    Функция формирования списка со сформированным
    текстом ответа.
    :param rows: данные по строке из датафрейма
    :return полученные данные
    """
    # Город
    result = {'city': rows[0]}
    result['id'] = rows[1]

    # Семейное положение (женат, холост/вдовец, вдова/разведена)
    marital_status = rows[17]

    if not marital_status:
        result['info'] = get_info_wife_block(rows).strip()
        result['education'] = get_press(rows).strip()
        result['socio_economic'] = get_socio_economic_block(rows).strip()
        result['family'] = ""
        result['sons'], result['daughters'] = get_children_block(rows)
    elif marital_status.strip() not in ['Холост', 'Вдовец']:
        result['info'] = get_info_block(rows).strip()
        result['education'] = get_education_block(rows).strip()
        result['socio_economic'] = get_socio_economic_block(rows).strip()
        result['family'] = get_family_block(rows).strip()
        result['sons'], result['daughters'] = get_children_block(rows)
    else:
        result['info'] = get_info_block(rows).strip()
        result['education'] = get_education_block(rows).strip()
        result['socio_economic'] = get_socio_economic_block(rows).strip()
        result['family'] = ""
        result['sons'], result['daughters'] = "", ""

    return result


def save_result(data: list):
    """
    Сохранение результатов в файл word
    :param data: список с прочитанными данными
    :return: None
    """
    doc = docx.Document()

    current_city = ''
    for i, item in enumerate(data):
        if i != 0 and item['city'] != current_city:
            doc.add_page_break()
        if item['city'] != current_city:
            title = doc.add_paragraph(item['city']+'\n')
            title.style = 'Heading 1'
            title.runs[0].underline = True

        if item['id']:
            par = doc.add_paragraph('№ анкеты: '+str(item['id']))
            par.runs[0].underline = True

        if item['info']:
            par = doc.add_paragraph(item['info'])
            par.alignment = 3

        if item['education']:
            par = doc.add_paragraph(item['education'])
            par.alignment = 3

        if item['socio_economic']:
            par = doc.add_paragraph(item['socio_economic'])
            par.alignment = 3

        if item['family']:
            par = doc.add_paragraph(item['family'])
            par.alignment = 3

        if item['sons']:
            par = doc.add_paragraph(item['sons'])
            par.alignment = 3

        if item['daughters']:
            par = doc.add_paragraph(item['daughters'])
            par.alignment = 3

        doc.add_paragraph('')
        current_city = item['city']

    doc.save('data/report.docx')


def main():
    data = []
    print('Считывание данных из xls-файла...')
    for i, rows in df.iterrows():
        data.append(create_text(rows=rows.fillna(0)))

    # Сохранение данных в нужном формате
    # JSON('data/output.json').write(data=data)
    print('Запись данных в doc-файл...')
    save_result(data)


if __name__ == '__main__':
    main()
